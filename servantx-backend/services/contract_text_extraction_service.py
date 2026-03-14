from pathlib import Path
from typing import List
import xml.etree.ElementTree as ET
import zipfile

from docx import Document as DocxDocument

from services.pdf_extraction_service import extract_text_from_pdf


WORD_XML_NAMES = (
    "word/document.xml",
    "word/header1.xml",
    "word/header2.xml",
    "word/header3.xml",
    "word/footer1.xml",
    "word/footer2.xml",
    "word/footer3.xml",
    "word/footnotes.xml",
    "word/endnotes.xml",
)


def _full_path_from_relative(file_path: str) -> Path:
    return Path("uploads") / file_path


def _extract_text_from_docx(relative_path: str) -> str:
    full_path = _full_path_from_relative(relative_path)
    if not full_path.exists():
        return f"File not found: {relative_path}"

    try:
        # Primary extractor: python-docx handles most real-world DOCX structures.
        doc = DocxDocument(str(full_path))
        paragraph_lines: List[str] = []
        for paragraph in doc.paragraphs:
            text = (paragraph.text or "").strip()
            if text:
                paragraph_lines.append(text)

        table_lines: List[str] = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = (cell.text or "").strip()
                    if text:
                        table_lines.append(text)

        primary_lines: List[str] = []
        seen = set()
        for line in paragraph_lines + table_lines:
            key = line.lower()
            if key in seen:
                continue
            seen.add(key)
            primary_lines.append(line)

        primary_output = "\n".join(primary_lines).strip()
        if primary_output:
            return primary_output
    except Exception:
        # Fall back to direct XML extraction below.
        pass

    try:
        collected_paragraphs: List[str] = []
        with zipfile.ZipFile(full_path, "r") as archive:
            names = set(archive.namelist())
            xml_names = [name for name in WORD_XML_NAMES if name in names]
            if not xml_names:
                return "Warning: DOCX structure missing Word document XML entries."

            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            for xml_name in xml_names:
                xml_bytes = archive.read(xml_name)
                root = ET.fromstring(xml_bytes)
                for paragraph in root.findall(".//w:p", ns):
                    runs = paragraph.findall(".//w:t", ns)
                    text = "".join(run.text or "" for run in runs).strip()
                    if text:
                        collected_paragraphs.append(text)

        output = "\n".join(collected_paragraphs).strip()
        if not output:
            return "Warning: No text could be extracted from DOCX."
        return output
    except Exception as exc:
        return f"Error extracting text from DOCX: {str(exc)}"


def extract_contract_text(relative_path: str, file_name: str = "") -> str:
    """
    Extract contract text from supported file types.
    """
    extension = (Path(file_name).suffix or Path(relative_path).suffix).lower()

    if extension == ".pdf":
        return extract_text_from_pdf(relative_path)
    if extension == ".docx":
        return _extract_text_from_docx(relative_path)
    if extension in (".txt", ".md"):
        full_path = _full_path_from_relative(relative_path)
        if not full_path.exists():
            return f"File not found: {relative_path}"
        try:
            return full_path.read_text(encoding="utf-8", errors="ignore").strip()
        except Exception as exc:
            return f"Error extracting text from text file: {str(exc)}"

    if extension == ".doc":
        return "Warning: Legacy .doc extraction is not currently supported. Please upload DOCX or PDF."

    return f"Warning: Unsupported contract file type: {extension or 'unknown'}"
